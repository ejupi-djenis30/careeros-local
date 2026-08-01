"""Bounded, database-free parsing for local Career Vault source documents."""

from __future__ import annotations

import hashlib
import json
import re
import zipfile
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path

from backend.career.schemas import SourceFactCandidate
from backend.core.config import settings


class SourceImportError(ValueError):
    pass


@dataclass(frozen=True)
class PreparedSourceDocument:
    original_name: str
    media_type: str
    document_type: str
    extracted_text: str = field(repr=False)
    sha256: str
    candidates: tuple[SourceFactCandidate, ...] = field(repr=False)
    data: bytes = field(repr=False)


def _bounded_source_text(text: str) -> str:
    normalized = text.strip()
    if len(normalized) > settings.SOURCE_IMPORT_MAX_EXTRACTED_CHARS:
        raise SourceImportError("Extracted source text exceeds the configured safety limit")
    return normalized


def _validate_docx_package(data: bytes) -> None:
    """Bound actual DOCX decompression before python-docx parses the package."""

    try:
        with zipfile.ZipFile(BytesIO(data)) as archive:
            members = archive.infolist()
            if len(members) > settings.SOURCE_IMPORT_MAX_ARCHIVE_MEMBERS:
                raise SourceImportError("The DOCX package contains too many members")

            seen: set[str] = set()
            total_bytes = 0
            has_document = False
            for info in members:
                canonical_name = info.filename.replace("\\", "/")
                folded_name = canonical_name.casefold()
                if (
                    not canonical_name
                    or "\x00" in canonical_name
                    or folded_name in seen
                    or info.flag_bits & 0x1
                ):
                    raise SourceImportError("The DOCX package structure is invalid")
                seen.add(folded_name)
                has_document = has_document or canonical_name == "word/document.xml"
                if info.is_dir():
                    continue
                if info.file_size < 0:
                    raise SourceImportError("The DOCX package size metadata is invalid")
                total_bytes += info.file_size
                if total_bytes > settings.SOURCE_IMPORT_MAX_UNCOMPRESSED_BYTES:
                    raise SourceImportError("The DOCX package expands beyond the safety limit")

                actual_member_bytes = 0
                with archive.open(info) as member:
                    while chunk := member.read(128 * 1024):
                        actual_member_bytes += len(chunk)
                        if actual_member_bytes > info.file_size:
                            raise SourceImportError("The DOCX package size metadata is invalid")
                if actual_member_bytes != info.file_size:
                    raise SourceImportError("The DOCX package size metadata is invalid")

            if not has_document:
                raise SourceImportError("The DOCX package has no document body")
    except SourceImportError:
        raise
    except (EOFError, OSError, RuntimeError, zipfile.BadZipFile) as exc:
        raise SourceImportError("The uploaded file is not a valid DOCX") from exc


def _document_type(filename: str, media_type: str) -> str:
    suffix = Path(filename).suffix.lower()
    media_type = media_type.split(";", 1)[0].strip().lower()
    if suffix in {".txt", ".md"} and media_type in {
        "text/plain",
        "text/markdown",
        "application/octet-stream",
    }:
        return "text"
    if suffix == ".pdf" and media_type in {"application/pdf", "application/octet-stream"}:
        return "pdf"
    if suffix == ".docx" and media_type in {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/octet-stream",
    }:
        return "docx"
    raise SourceImportError("Supported source formats are TXT, Markdown, PDF and DOCX")


def extract_text(data: bytes, document_type: str) -> str:
    if document_type == "text":
        try:
            return _bounded_source_text(data.decode("utf-8-sig").replace("\x00", ""))
        except UnicodeDecodeError as exc:
            raise SourceImportError("Text documents must use UTF-8 encoding") from exc
    if document_type == "pdf":
        if not data.startswith(b"%PDF"):
            raise SourceImportError("The uploaded file is not a valid PDF")
        from pypdf import PdfReader

        try:
            pdf_document = PdfReader(BytesIO(data))
            if len(pdf_document.pages) > settings.SOURCE_IMPORT_MAX_PAGES:
                raise SourceImportError("The PDF exceeds the configured page limit")
            parts: list[str] = []
            extracted_characters = 0
            for page in pdf_document.pages:
                page_text = page.extract_text() or ""
                extracted_characters += len(page_text)
                if extracted_characters > settings.SOURCE_IMPORT_MAX_EXTRACTED_CHARS:
                    raise SourceImportError(
                        "Extracted source text exceeds the configured safety limit"
                    )
                parts.append(page_text)
            return _bounded_source_text("\n".join(parts))
        except SourceImportError:
            raise
        except Exception as exc:
            raise SourceImportError("Unable to read the PDF") from exc
    if not data.startswith(b"PK"):
        raise SourceImportError("The uploaded file is not a valid DOCX")
    _validate_docx_package(data)
    try:
        from docx import Document

        word_document = Document(BytesIO(data))
        return _bounded_source_text(
            "\n".join(paragraph.text for paragraph in word_document.paragraphs)
        )
    except SourceImportError:
        raise
    except Exception as exc:
        raise SourceImportError("Unable to read the DOCX") from exc


_SKILL_LIST = re.compile(
    r"^(?:skills?|competenze|technologies|tecnologie)\s*[:\-]\s*(.+)$",
    re.IGNORECASE,
)


def _candidate_id(locator: str, fact_type: str, payload: dict[str, object]) -> str:
    canonical = json.dumps(payload, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
    return hashlib.sha256(f"{locator}\0{fact_type}\0{canonical}".encode()).hexdigest()


def _candidate_blocks(text: str) -> list[tuple[str, str]]:
    paragraphs = [item.strip() for item in re.split(r"\n\s*\n+", text) if item.strip()]
    if len(paragraphs) <= 1:
        paragraphs = [item.strip() for item in text.splitlines() if item.strip()]
    result: list[tuple[str, str]] = []
    seen: set[str] = set()
    for index, raw in enumerate(paragraphs, start=1):
        normalized = re.sub(r"\s+", " ", raw).strip()
        key = normalized.casefold()
        if not normalized or key in seen:
            continue
        seen.add(key)
        result.append((f"paragraph:{index}", normalized[:5000]))
    return result


def fact_candidates(text: str) -> list[SourceFactCandidate]:
    candidates: list[SourceFactCandidate] = []
    for locator, block in _candidate_blocks(text):
        skill_match = _SKILL_LIST.match(block)
        if skill_match:
            skills = [
                value.strip(" .")
                for value in re.split(r"[,;|]", skill_match.group(1))
                if value.strip(" .")
            ]
            for item_index, skill in enumerate(skills[:12], start=1):
                payload: dict[str, object] = {"name": skill[:160], "level": "working"}
                skill_locator = f"{locator}:skill:{item_index}"
                candidates.append(
                    SourceFactCandidate(
                        candidate_id=_candidate_id(skill_locator, "skill", payload),
                        fact_type="skill",
                        payload=payload,
                        source_locator=skill_locator,
                        confidence=0.82,
                        excerpt=block[:1000],
                    )
                )
            continue
        if len(block) < 20 or block.endswith(":"):
            continue
        title = re.split(r"[.!?]", block, maxsplit=1)[0].strip(" -•\t")[:240]
        if not title:
            continue
        payload = {"title": title, "description": block[:5000]}
        candidates.append(
            SourceFactCandidate(
                candidate_id=_candidate_id(locator, "achievement", payload),
                fact_type="achievement",
                payload=payload,
                source_locator=locator,
                confidence=0.58,
                excerpt=block[:1000],
            )
        )
        if len(candidates) >= 24:
            break
    return candidates[:24]


def prepare_source_document(
    *,
    filename: str,
    media_type: str,
    data: bytes,
) -> PreparedSourceDocument:
    if not data:
        raise SourceImportError("The source document is empty")
    if len(data) > settings.MAX_UPLOAD_FILE_SIZE:
        raise SourceImportError("The source document exceeds the configured size limit")
    safe_name = str(filename or "source").replace("\\", "/").rsplit("/", 1)[-1]
    safe_name = (
        "".join(
            character for character in safe_name if ord(character) >= 32 and ord(character) != 127
        ).strip(" .")[:255]
        or "source"
    )
    normalized_media_type = (media_type or "application/octet-stream").split(";", 1)[0]
    document_type = _document_type(safe_name, normalized_media_type)
    extracted_text = extract_text(data, document_type)
    return PreparedSourceDocument(
        original_name=safe_name,
        media_type=normalized_media_type,
        document_type=document_type,
        extracted_text=extracted_text,
        sha256=hashlib.sha256(data).hexdigest(),
        candidates=tuple(fact_candidates(extracted_text)),
        data=data,
    )
