import hashlib
import json
import re
from io import BytesIO
from pathlib import Path

from sqlalchemy.orm import Session

from backend.career.models import CareerAsset, SourceDocument
from backend.career.repository import CareerProfileRepository
from backend.career.schemas import SourceDocumentResponse, SourceFactCandidate
from backend.core.config import settings
from backend.storage.atomic import atomic_write, resolve_data_path


class SourceImportError(ValueError):
    pass


def _document_type(filename: str, media_type: str) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix in {".txt", ".md"} and media_type in {
        "text/plain",
        "text/markdown",
        "application/octet-stream",
    }:
        return "text"
    if suffix == ".pdf" and media_type in {"application/pdf", "application/octet-stream"}:
        return "pdf"
    if suffix == ".docx" and media_type in {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/octet-stream",
    }:
        return "docx"
    raise SourceImportError("Supported source formats are TXT, Markdown, PDF and DOCX")


def _extract_text(data: bytes, document_type: str) -> str:
    if document_type == "text":
        try:
            return data.decode("utf-8-sig").replace("\x00", "").strip()
        except UnicodeDecodeError as exc:
            raise SourceImportError("Text documents must use UTF-8 encoding") from exc
    if document_type == "pdf":
        if not data.startswith(b"%PDF"):
            raise SourceImportError("The uploaded file is not a valid PDF")
        from pypdf import PdfReader

        try:
            pdf_document = PdfReader(BytesIO(data))
            return "\n".join(
                page.extract_text() or "" for page in pdf_document.pages
            ).strip()
        except Exception as exc:
            raise SourceImportError("Unable to read the PDF") from exc
    if not data.startswith(b"PK"):
        raise SourceImportError("The uploaded file is not a valid DOCX")
    try:
        from docx import Document

        word_document = Document(BytesIO(data))
        return "\n".join(paragraph.text for paragraph in word_document.paragraphs).strip()
    except Exception as exc:
        raise SourceImportError("Unable to read the DOCX") from exc


_SKILL_LIST = re.compile(
    r"^(?:skills?|competenze|technologies|tecnologie)\s*[:\-]\s*(.+)$",
    re.IGNORECASE,
)


def _candidate_id(locator: str, fact_type: str, payload: dict[str, object]) -> str:
    canonical = json.dumps(payload, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
    return hashlib.sha256(f"{locator}\0{fact_type}\0{canonical}".encode()).hexdigest()


def _candidate_blocks(text: str) -> list[tuple[str, str]]:
    paragraphs = [item.strip() for item in re.split(r"\n\s*\n+", text) if item.strip()]
    if len(paragraphs) <= 1:
        paragraphs = [item.strip() for item in text.splitlines() if item.strip()]
    result: list[tuple[str, str]] = []
    seen: set[str] = set()
    for index, raw in enumerate(paragraphs, start=1):
        normalized = re.sub(r"\s+", " ", raw).strip()
        key = normalized.casefold()
        if not normalized or key in seen:
            continue
        seen.add(key)
        result.append((f"paragraph:{index}", normalized[:5000]))
    return result


def fact_candidates(text: str) -> list[SourceFactCandidate]:
    candidates: list[SourceFactCandidate] = []
    for locator, block in _candidate_blocks(text):
        skill_match = _SKILL_LIST.match(block)
        if skill_match:
            skills = [
                value.strip(" .")
                for value in re.split(r"[,;|]", skill_match.group(1))
                if value.strip(" .")
            ]
            for item_index, skill in enumerate(skills[:12], start=1):
                payload: dict[str, object] = {"name": skill[:160], "level": "working"}
                skill_locator = f"{locator}:skill:{item_index}"
                candidates.append(
                    SourceFactCandidate(
                        candidate_id=_candidate_id(skill_locator, "skill", payload),
                        fact_type="skill",
                        payload=payload,
                        source_locator=skill_locator,
                        confidence=0.82,
                        excerpt=block[:1000],
                    )
                )
            continue
        if len(block) < 20 or block.endswith(":"):
            continue
        title = re.split(r"[.!?]", block, maxsplit=1)[0].strip(" -•\t")[:240]
        if not title:
            continue
        payload = {"title": title, "description": block[:5000]}
        candidates.append(
            SourceFactCandidate(
                candidate_id=_candidate_id(locator, "achievement", payload),
                fact_type="achievement",
                payload=payload,
                source_locator=locator,
                confidence=0.58,
                excerpt=block[:1000],
            )
        )
        if len(candidates) >= 24:
            break
    return candidates[:24]


def _response(source: SourceDocument) -> SourceDocumentResponse:
    return SourceDocumentResponse(
        id=source.id,
        asset_id=source.asset.id,
        original_name=source.asset.original_name,
        media_type=source.asset.media_type,
        sha256=source.asset.sha256,
        byte_size=source.asset.byte_size,
        document_type=source.document_type,
        extracted_characters=len(source.extracted_text),
        text_preview=source.extracted_text[:4000],
        candidates=fact_candidates(source.extracted_text),
        created_at=source.created_at,
    )


def import_source_document(
    db: Session,
    *,
    user_id: int,
    filename: str,
    media_type: str,
    data: bytes,
) -> SourceDocumentResponse:
    profile = CareerProfileRepository(db).get_by_user(user_id)
    if profile is None:
        raise SourceImportError("Create the career profile before importing source documents")
    if not data:
        raise SourceImportError("The source document is empty")
    if len(data) > settings.MAX_UPLOAD_FILE_SIZE:
        raise SourceImportError("The source document exceeds the configured size limit")

    safe_name = Path(filename or "source").name[:255]
    kind = _document_type(safe_name, media_type or "application/octet-stream")
    extracted = _extract_text(data, kind)
    digest = hashlib.sha256(data).hexdigest()
    existing = (
        db.query(SourceDocument)
        .join(CareerAsset, SourceDocument.asset_id == CareerAsset.id)
        .filter(
            CareerAsset.profile_id == profile.id,
            CareerAsset.sha256 == digest,
            CareerAsset.kind == "source_document",
        )
        .first()
    )
    if existing:
        return _response(existing)

    relative_path = Path("assets") / digest[:2] / digest
    absolute_path, created_file = atomic_write(relative_path, data)
    try:
        asset = CareerAsset(
            profile_id=profile.id,
            kind="source_document",
            original_name=safe_name,
            media_type=media_type or "application/octet-stream",
            sha256=digest,
            byte_size=len(data),
            storage_path=relative_path.as_posix(),
            normalized=False,
        )
        db.add(asset)
        db.flush()
        source = SourceDocument(
            profile_id=profile.id,
            asset_id=asset.id,
            document_type=kind,
            extracted_text=extracted,
            extracted_text_sha256=hashlib.sha256(extracted.encode("utf-8")).hexdigest(),
        )
        db.add(source)
        db.commit()
        db.refresh(source)
        return _response(source)
    except Exception:
        db.rollback()
        if created_file:
            resolve_data_path(relative_path).unlink(missing_ok=True)
        raise
