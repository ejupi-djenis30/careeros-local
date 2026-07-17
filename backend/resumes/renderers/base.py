from html import escape
from io import BytesIO

from docx import Document as create_document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer

from backend.resumes.content import ResumeContent, build_content

PDF_MEDIA_TYPE = "application/pdf"
DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _paragraph_text(value: str) -> str:
    return escape(value).replace("\n", "<br/>")


def render_pdf(snapshot: dict, *, photo: bytes | None = None) -> bytes:
    content = build_content(snapshot)
    style = content.style
    margin = float(style.get("margin_mm", 17))
    base_size = float(style.get("base_font_size", 9))
    line_height = float(style.get("line_height", 1.3))
    spacing = float(style.get("section_spacing", 8))
    accent = str(style.get("accent_color", "#111827"))
    requested_font = str(style.get("font_family", "Helvetica"))
    font_names = {
        "Helvetica": ("Helvetica", "Helvetica-Bold", "Helvetica-Oblique"),
        "Arial": ("Helvetica", "Helvetica-Bold", "Helvetica-Oblique"),
        "Georgia": ("Times-Roman", "Times-Bold", "Times-Italic"),
    }
    normal_font, bold_font, italic_font = font_names.get(
        requested_font, font_names["Helvetica"]
    )
    output = BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=margin * mm,
        leftMargin=margin * mm,
        topMargin=margin * mm,
        bottomMargin=margin * mm,
        title=content.display_name,
        author="CareerOS Local",
        subject="Resume",
    )
    base_styles = getSampleStyleSheet()
    name_style = ParagraphStyle(
        "ResumeName",
        parent=base_styles["Title"],
        fontName=bold_font,
        fontSize=20,
        leading=23,
        alignment=TA_CENTER,
        spaceAfter=3,
    )
    headline_style = ParagraphStyle(
        "ResumeHeadline",
        parent=base_styles["Normal"],
        fontName=normal_font,
        fontSize=base_size + 2,
        leading=(base_size + 2) * line_height,
        alignment=TA_CENTER,
        textColor=HexColor("#30343B"),
        spaceAfter=3,
    )
    contact_style = ParagraphStyle(
        "ResumeContact",
        parent=base_styles["Normal"],
        fontName=normal_font,
        fontSize=8.5,
        leading=11,
        alignment=TA_CENTER,
        textColor=HexColor("#4B5563"),
        spaceAfter=8,
    )
    section_style = ParagraphStyle(
        "ResumeSection",
        parent=base_styles["Heading2"],
        fontName=bold_font,
        fontSize=base_size + 1,
        leading=(base_size + 1) * line_height,
        textColor=HexColor(accent),
        spaceBefore=spacing,
        spaceAfter=4,
        borderWidth=0,
        borderPadding=0,
    )
    entry_title_style = ParagraphStyle(
        "ResumeEntryTitle",
        parent=base_styles["Normal"],
        fontName=bold_font,
        fontSize=base_size + 0.5,
        leading=(base_size + 0.5) * line_height,
        spaceBefore=3,
        spaceAfter=1,
    )
    meta_style = ParagraphStyle(
        "ResumeMeta",
        parent=base_styles["Normal"],
        fontName=italic_font,
        fontSize=8.5,
        leading=11,
        textColor=HexColor("#374151"),
        spaceAfter=2,
    )
    body_style = ParagraphStyle(
        "ResumeBody",
        parent=base_styles["Normal"],
        fontName=normal_font,
        fontSize=base_size,
        leading=base_size * line_height,
        spaceAfter=2,
    )
    bullet_style = ParagraphStyle(
        "ResumeBullet",
        parent=body_style,
        leftIndent=10,
        firstLineIndent=-7,
        bulletIndent=2,
    )

    story = []
    if photo:
        image = Image(BytesIO(photo), width=28 * mm, height=28 * mm)
        image.hAlign = "CENTER"
        story.extend([image, Spacer(1, 4)])
    story.append(Paragraph(_paragraph_text(content.display_name), name_style))
    if content.headline:
        story.append(Paragraph(_paragraph_text(content.headline), headline_style))
    if content.contact_line:
        story.append(Paragraph(_paragraph_text(content.contact_line), contact_style))
    if content.summary:
        story.append(Paragraph(_paragraph_text(content.summary_heading), section_style))
        story.append(Paragraph(_paragraph_text(content.summary), body_style))
    for section in content.sections:
        if section.page_break_before:
            story.append(PageBreak())
        story.append(Paragraph(_paragraph_text(section.heading), section_style))
        for entry in section.entries:
            story.append(Paragraph(_paragraph_text(entry.title), entry_title_style))
            metadata = " | ".join(value for value in (entry.subtitle, entry.date_range) if value)
            if metadata:
                story.append(Paragraph(_paragraph_text(metadata), meta_style))
            if entry.description:
                story.append(Paragraph(_paragraph_text(entry.description), body_style))
            for bullet in entry.bullets:
                story.append(Paragraph(f"&bull;&nbsp;{_paragraph_text(bullet)}", bullet_style))
    document.build(story)
    return output.getvalue()


def _configure_docx(document: DocxDocument, title: str, style: dict) -> None:
    section = document.sections[0]
    margin_inches = float(style.get("margin_mm", 16.5)) / 25.4
    section.top_margin = Inches(margin_inches)
    section.bottom_margin = Inches(margin_inches)
    section.left_margin = Inches(margin_inches)
    section.right_margin = Inches(margin_inches)
    normal = document.styles["Normal"]
    normal.font.name = str(style.get("font_family", "Arial"))
    normal.font.size = Pt(float(style.get("base_font_size", 9.5)))
    normal.paragraph_format.space_after = Pt(2)
    properties = document.core_properties
    properties.title = title
    properties.author = "CareerOS Local"
    properties.subject = "Resume"
    properties.keywords = ""
    properties.comments = "Generated locally"


def _add_docx_heading(document: DocxDocument, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(10)


def _add_docx_entry(document: DocxDocument, entry) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(entry.title)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    metadata = " | ".join(value for value in (entry.subtitle, entry.date_range) if value)
    if metadata:
        paragraph = document.add_paragraph(metadata)
        paragraph.paragraph_format.space_after = Pt(1)
        for run in paragraph.runs:
            run.italic = True
            run.font.size = Pt(8.5)
    if entry.description:
        document.add_paragraph(entry.description)
    for bullet in entry.bullets:
        paragraph = document.add_paragraph(bullet, style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(1)


def render_docx(snapshot: dict, *, photo: bytes | None = None) -> bytes:
    content: ResumeContent = build_content(snapshot)
    document = create_document()
    _configure_docx(document, content.display_name, content.style)
    if photo:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(BytesIO(photo), width=Inches(1.05))
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(1)
    run = paragraph.add_run(content.display_name)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(20)
    if content.headline:
        paragraph = document.add_paragraph(content.headline)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(1)
    if content.contact_line:
        paragraph = document.add_paragraph(content.contact_line)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(5)
        for run in paragraph.runs:
            run.font.size = Pt(8.5)
    if content.summary:
        _add_docx_heading(document, content.summary_heading)
        document.add_paragraph(content.summary)
    for section in content.sections:
        if section.page_break_before:
            document.add_page_break()
        _add_docx_heading(document, section.heading)
        for entry in section.entries:
            _add_docx_entry(document, entry)
    output = BytesIO()
    document.save(output)
    return output.getvalue()
