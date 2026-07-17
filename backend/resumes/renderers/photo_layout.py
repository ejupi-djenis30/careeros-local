from html import escape
from io import BytesIO

from docx import Document as create_document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    BalancedColumns,
    Image,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)

from backend.resumes.content import ResumeContent, build_content
from backend.resumes.renderers.base import _add_docx_entry, _add_docx_heading, _configure_docx


def _text(value: str) -> str:
    return escape(value).replace("\n", "<br/>")


def _pdf_styles(content: ResumeContent) -> dict[str, ParagraphStyle]:
    style = content.style
    base_size = float(style.get("base_font_size", 9))
    line_height = float(style.get("line_height", 1.3))
    spacing = float(style.get("section_spacing", 8))
    font_names = {
        "Helvetica": ("Helvetica", "Helvetica-Bold", "Helvetica-Oblique"),
        "Arial": ("Helvetica", "Helvetica-Bold", "Helvetica-Oblique"),
        "Georgia": ("Times-Roman", "Times-Bold", "Times-Italic"),
    }
    normal, bold, italic = font_names.get(
        str(style.get("font_family", "Helvetica")), font_names["Helvetica"]
    )
    base = getSampleStyleSheet()
    body = ParagraphStyle(
        "PhotoBody",
        parent=base["Normal"],
        fontName=normal,
        fontSize=base_size,
        leading=base_size * line_height,
        spaceAfter=2,
    )
    return {
        "name": ParagraphStyle(
            "PhotoName",
            parent=base["Title"],
            fontName=bold,
            fontSize=20,
            leading=23,
            alignment=TA_CENTER,
            spaceAfter=3,
        ),
        "headline": ParagraphStyle(
            "PhotoHeadline",
            parent=body,
            fontSize=base_size + 2,
            leading=(base_size + 2) * line_height,
            alignment=TA_CENTER,
            textColor=HexColor("#30343B"),
        ),
        "contact": ParagraphStyle(
            "PhotoContact",
            parent=body,
            fontSize=8.5,
            leading=11,
            alignment=TA_CENTER,
            textColor=HexColor("#4B5563"),
            spaceAfter=8,
        ),
        "section": ParagraphStyle(
            "PhotoSection",
            parent=base["Heading2"],
            fontName=bold,
            fontSize=base_size + 1,
            leading=(base_size + 1) * line_height,
            textColor=HexColor(str(style.get("accent_color", "#111827"))),
            spaceBefore=spacing,
            spaceAfter=4,
        ),
        "title": ParagraphStyle(
            "PhotoEntryTitle",
            parent=body,
            fontName=bold,
            fontSize=base_size + 0.5,
            leading=(base_size + 0.5) * line_height,
            spaceBefore=3,
            spaceAfter=1,
        ),
        "meta": ParagraphStyle(
            "PhotoMeta",
            parent=body,
            fontName=italic,
            fontSize=8.5,
            leading=11,
            textColor=HexColor("#374151"),
        ),
        "body": body,
        "bullet": ParagraphStyle(
            "PhotoBullet",
            parent=body,
            leftIndent=10,
            firstLineIndent=-7,
            bulletIndent=2,
        ),
    }


def _section_flowables(section, styles: dict[str, ParagraphStyle]) -> list:
    result = [Paragraph(_text(section.heading), styles["section"])]
    for entry in section.entries:
        result.append(Paragraph(_text(entry.title), styles["title"]))
        metadata = " | ".join(value for value in (entry.subtitle, entry.date_range) if value)
        if metadata:
            result.append(Paragraph(_text(metadata), styles["meta"]))
        if entry.description:
            result.append(Paragraph(_text(entry.description), styles["body"]))
        result.extend(
            Paragraph(f"&bull;&nbsp;{_text(bullet)}", styles["bullet"])
            for bullet in entry.bullets
        )
    return result


def render_two_column_pdf(snapshot: dict, photo: bytes | None) -> bytes:
    content = build_content(snapshot)
    margin = float(content.style.get("margin_mm", 17))
    output = BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=margin * mm,
        leftMargin=margin * mm,
        topMargin=margin * mm,
        bottomMargin=margin * mm,
        title=content.display_name,
        author="CareerOS Local",
        subject="Resume",
    )
    styles = _pdf_styles(content)
    story = []
    if photo:
        image = Image(BytesIO(photo), width=28 * mm, height=28 * mm)
        image.hAlign = "CENTER"
        story.extend([image, Spacer(1, 4)])
    story.append(Paragraph(_text(content.display_name), styles["name"]))
    if content.headline:
        story.append(Paragraph(_text(content.headline), styles["headline"]))
    if content.contact_line:
        story.append(Paragraph(_text(content.contact_line), styles["contact"]))
    if content.summary:
        story.append(Paragraph(_text(content.summary_heading), styles["section"]))
        story.append(Paragraph(_text(content.summary), styles["body"]))

    segments: list[list] = [[]]
    for section in content.sections:
        if section.page_break_before and segments[-1]:
            segments.append([])
        segments[-1].extend(_section_flowables(section, styles))
    for index, segment in enumerate(filter(None, segments)):
        if index:
            story.append(PageBreak())
        story.append(BalancedColumns(segment, nCols=2, innerPadding=8 * mm, endSlack=0.05))
    document.build(story)
    return output.getvalue()


def _add_header(document, content: ResumeContent, photo: bytes | None) -> None:
    if photo:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(BytesIO(photo), width=Inches(1.05))
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(content.display_name)
    run.bold = True
    run.font.size = Pt(20)
    if content.headline:
        paragraph = document.add_paragraph(content.headline)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if content.contact_line:
        paragraph = document.add_paragraph(content.contact_line)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.size = Pt(8.5)
    if content.summary:
        _add_docx_heading(document, content.summary_heading)
        document.add_paragraph(content.summary)


def render_two_column_docx(snapshot: dict, photo: bytes | None) -> bytes:
    content = build_content(snapshot)
    document = create_document()
    _configure_docx(document, content.display_name, content.style)
    _add_header(document, content, photo)
    section = document.add_section(WD_SECTION.CONTINUOUS)
    columns = section._sectPr.xpath("./w:cols")[0]
    columns.set(qn("w:num"), "2")
    columns.set(qn("w:space"), "360")
    for item in content.sections:
        if item.page_break_before:
            document.add_page_break()
        _add_docx_heading(document, item.heading)
        for entry in item.entries:
            _add_docx_entry(document, entry)
    output = BytesIO()
    document.save(output)
    return output.getvalue()
