import zipfile
from datetime import datetime, timezone
from io import BytesIO

import pytest
from docx import Document
from pypdf import PdfReader

from backend.core.config import settings
from backend.resumes.quality import ResumeQualityError, validate_resume_artifacts
from backend.resumes.renderers.ats import render_ats_docx, render_ats_pdf


def _snapshot(entry_count: int = 2) -> dict:
    blocks = [
        {
            "id": f"fact-{index}",
            "kind": "fact",
            "fact_ids": [f"10000000-0000-4000-8000-{index:012d}"],
            "visible": True,
            "content": {
                "title": f"Role {index}",
                "subtitle": "Local Systems",
                "date_range": "2020 – 2026",
                "description": "Built dependable private systems. " * 8,
                "bullets": ["Reduced lead time by 40 percent."],
            },
            "manual_fields": [],
            "layout": {"spacing_before_pt": 0, "keep_together": True},
        }
        for index in range(entry_count)
    ]
    return {
        "profile": {
            "display_name": "Mira Vale",
            "headline": "Principal Engineer",
            "email": "mira@example.test",
        },
        "resume": {
            "title": "Private systems CV",
            "template_kind": "ats",
            "section_config": {"order": ["experience"]},
            "content_overrides": {},
            "canvas_document": {
                "schema_version": 2,
                "style": {
                    "font_family": "Helvetica",
                    "base_font_size": 10,
                    "line_height": 1.3,
                    "section_spacing": 10,
                    "margin_mm": 18,
                    "accent_color": "#243B53",
                    "columns": 1,
                },
                "sections": [
                    {
                        "id": "identity",
                        "kind": "identity",
                        "title": "IDENTITY",
                        "visible": True,
                        "page_break_before": False,
                        "blocks": [
                            {
                                "id": "identity-main",
                                "kind": "identity",
                                "fact_ids": [],
                                "visible": True,
                                "content": {
                                    "title": "Mira Vale",
                                    "subtitle": "Principal Engineer",
                                    "date_range": "",
                                    "description": "mira@example.test",
                                    "bullets": [],
                                },
                                "manual_fields": [],
                                "layout": {"spacing_before_pt": 0, "keep_together": True},
                            }
                        ],
                    },
                    {
                        "id": "experience",
                        "kind": "experience",
                        "title": "EXPERIENCE",
                        "visible": True,
                        "page_break_before": False,
                        "blocks": blocks,
                    },
                ],
            },
        },
        "facts": [],
    }


def test_ats_renderers_preserve_text_order_and_local_metadata():
    snapshot = _snapshot()
    pdf = render_ats_pdf(snapshot)
    docx = render_ats_docx(snapshot)
    report = validate_resume_artifacts(
        pdf=pdf,
        docx=docx,
        required_headings=["EXPERIENCE"],
        required_text=["Mira Vale", "Role 0", "Role 1"],
        template_kind="ats",
        expect_photo=False,
    )

    pdf_document = PdfReader(BytesIO(pdf))
    text = "\n".join(page.extract_text() or "" for page in pdf_document.pages)
    assert pdf_document.metadata is not None
    assert pdf_document.metadata.author == "CareerOS Local"
    assert pdf_document.metadata.creator == "CareerOS Local"
    assert pdf_document.metadata.producer == "CareerOS Local"
    assert pdf_document.metadata.title == "Mira Vale"
    word = Document(BytesIO(docx))
    word_text = "\n".join(paragraph.text for paragraph in word.paragraphs)
    assert word.core_properties.author == "CareerOS Local"
    assert word.core_properties.comments == "Generated locally"
    assert word.core_properties.last_modified_by == "CareerOS Local"
    assert text.index("EXPERIENCE") < text.index("Role 0") < text.index("Role 1")
    assert word_text.index("EXPERIENCE") < word_text.index("Role 0") < word_text.index("Role 1")
    assert report["text_order_verified"] is True
    assert report["metadata_sanitized"] is True


def test_renderers_are_byte_stable_and_strip_host_timestamps():
    snapshot = _snapshot()

    first_pdf = render_ats_pdf(snapshot)
    first_docx = render_ats_docx(snapshot)
    assert first_pdf == render_ats_pdf(snapshot)
    assert first_docx == render_ats_docx(snapshot)

    pdf = PdfReader(BytesIO(first_pdf))
    assert pdf.metadata is not None
    assert pdf.metadata.creation_date == datetime(2000, 1, 1, tzinfo=timezone.utc)
    assert pdf.metadata.modification_date == datetime(2000, 1, 1, tzinfo=timezone.utc)
    word = Document(BytesIO(first_docx))
    assert word.core_properties.created == datetime(2000, 1, 1, tzinfo=timezone.utc)
    assert word.core_properties.modified == datetime(2000, 1, 1, tzinfo=timezone.utc)
    with zipfile.ZipFile(BytesIO(first_docx)) as archive:
        assert {entry.date_time for entry in archive.infolist()} == {(1980, 1, 1, 0, 0, 0)}
        assert {entry.extra for entry in archive.infolist()} == {b""}
        assert {entry.comment for entry in archive.infolist()} == {b""}


def test_quality_gate_rejects_artifact_size_before_parsing(monkeypatch):
    monkeypatch.setattr("backend.resumes.quality.MAX_RESUME_ARTIFACT_BYTES", 8)

    with pytest.raises(ResumeQualityError, match="cannot exceed"):
        validate_resume_artifacts(
            pdf=b"x" * 9,
            docx=b"y",
            required_headings=[],
            required_text=[],
            template_kind="ats",
            expect_photo=False,
        )


def test_quality_gate_preflights_docx_expansion(monkeypatch):
    output = BytesIO()
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", b"types")
        archive.writestr("word/document.xml", b"x" * 128)
    monkeypatch.setattr("backend.resumes.quality.MAX_RESUME_DOCX_UNCOMPRESSED_BYTES", 64)

    with pytest.raises(ResumeQualityError, match="expands beyond"):
        validate_resume_artifacts(
            pdf=render_ats_pdf(_snapshot()),
            docx=output.getvalue(),
            required_headings=[],
            required_text=[],
            template_kind="ats",
            expect_photo=False,
        )


def test_quality_gate_rejects_page_overflow(monkeypatch):
    monkeypatch.setattr(settings, "RESUME_MAX_PAGES", 1)
    snapshot = _snapshot(entry_count=24)
    pdf = render_ats_pdf(snapshot)
    docx = render_ats_docx(snapshot)
    with pytest.raises(ResumeQualityError, match="page"):
        validate_resume_artifacts(
            pdf=pdf,
            docx=docx,
            required_headings=["EXPERIENCE"],
            required_text=["Mira Vale"],
            template_kind="ats",
            expect_photo=False,
        )
